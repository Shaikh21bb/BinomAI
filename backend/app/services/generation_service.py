import json
import re
import time
import uuid
import structlog
from typing import Any, Dict, List, Optional, Tuple

import httpx
from fastapi import HTTPException
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.ai.llm_client import call_llm
from app.db.models.generated_document import GeneratedDocument
from app.schemas.generated_document import GenerateRequest, GeneratedContent

logger = structlog.get_logger(__name__)

re_fullmatch = re.fullmatch

GENERATION_SYSTEM_PROMPT = """Вы — ИИ-эксперт BINOM AI по подготовке тендерных документов для государственных и частных закупок Казахстана.
Сгенерируйте деловой документ строго на основе предоставленных данных: анализа технического задания, уточнённых данных компании из диалога и профиля компании.
Правила:
- Не выдумывайте факты, которых нет в данных. Если данных не хватает, сформулируйте нейтрально или укажите "будет подтверждено".
- Числа, суммы, сроки берите только из данных.
- Документ должен быть официальным, лаконичным и готовым к подаче.
- Если задание требует двуязычного документа (русский + казахский) — сначала идёт ПОЛНЫЙ документ на русском, затем ПОЛНЫЙ документ на казахском, без сокращений.
- Разрешены markdown-таблицы (строки вида "| Колонка 1 | Колонка 2 |" и разделитель "|---|---|").
- Верните строго JSON: {"title": "...", "content_md": "..."} где content_md — полный текст документа в Markdown (заголовки, списки, таблицы, абзацы).
"""

DOC_TYPE_LABELS = {
    "commercial_proposal": "Коммерческое предложение",
    "tech_spec": "Техническая спецификация",
    "cover_letter": "Сопроводительное письмо",
}

DOC_TYPE_PROMPTS = {
    "commercial_proposal": """СОСТАВЬТЕ КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:
- Обращение к заказчику
- Стоимость работ (только из данных диалога/анализа, с НДС)
- Условия оплаты (аванс, этапы — если указаны)
- Сроки выполнения работ
- Подпись: от имени компании""",
    "tech_spec": """СОСТАВЬТЕ ПОЛНУЮ ТЕХНИЧЕСКУЮ СПЕЦИФИКАЦИЮ (ТЕХНИЧЕСКОЕ ЗАДАНИЕ) в структуре официальной проектной документации РК:
ЧАСТЬ 1 — НА РУССКОМ ЯЗЫКЕ (полный документ):
1. Общие данные: наименование объекта, заказчик, основание для выполнения работ
2. Общие сведения об объекте: адрес/расположение, характеристика существующего здания/сооружения (конструктивные элементы)
3. Общие указания: условия площадки, нормативные документы (СП РК, СН РК, ГОСТ), производство работ, охрана труда
4. Состав и перечень работ: каждый пункт работ отдельным пунктом списка (как ведомость работ с количествами, если есть данные)
5. Требования к материалам и оборудованию: наименование, характеристика, соответствие ГОСТ/СП РК
6. Спецификация элементов/материалов: МАРКДАУН-ТАБЛИЦА с колонками (№, Наименование, Характеристика, Ед. изм., Количество)
7. Сроки выполнения и условия производства работ (из анализа/диалога)
8. Технико-экономические показатели (если есть данные)
9. Гарантийные обязательства

ЧАСТЬ 2 — ТОТ ЖЕ ПОЛНЫЙ ДОКУМЕНТ НА КАЗАХСКОМ ЯЗЫКЕ (весь документ заново, разделы 1-9, официальный казахский деловой стиль).
ВСЕ разделы и таблицы должны идти и на русском, и на казахском языке. Используйте markdown-таблицы для спецификаций.""",
    "cover_letter": """СОСТАВЬТЕ СОПРОВОДИТЕЛЬНОЕ ПИСЬМО:
- Обращение к заказчику с указанием предмета тендера
- Перечень прилагаемых документов (коммерческое предложение, техническая спецификация, документы компании)
- Готовность предоставить дополнительные материалы
- Подпись: от имени компании""",
}


def _md_to_html(md: str) -> str:
    """Minimal markdown -> HTML conversion for LibreOffice (headings, lists, tables, paragraphs)."""
    import html as html_mod
    from markdown_it import MarkdownIt

    try:
        return MarkdownIt("commonmark").enable("table").render(md)
    except Exception:
        pass

    lines = []
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            lines.append(f"<h3>{html_mod.escape(line[4:])}</h3>")
        elif line.startswith("## "):
            lines.append(f"<h2>{html_mod.escape(line[3:])}</h2>")
        elif line.startswith("# "):
            lines.append(f"<h1>{html_mod.escape(line[2:])}</h1>")
        elif line.startswith("- ") or line.startswith("* "):
            lines.append(f"<li>{html_mod.escape(line[2:])}</li>")
        else:
            lines.append(f"<p>{html_mod.escape(line)}</p>")
    return "\n".join(lines)


def _full_html(title: str, body_html: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8">
<style>
body {{ font-family: 'Arial', sans-serif; font-size: 11pt; line-height: 1.5; color: #111; }}
h1 {{ font-size: 15pt; text-align: center; }}
h2 {{ font-size: 13pt; }}
h3 {{ font-size: 12pt; }}
p {{ margin: 6pt 0; }}
li {{ margin: 2pt 0; }}
</style></head><body><h1>{title}</h1>{body_html}</body></html>"""


class GenerationService:
    @staticmethod
    async def _load_analysis_context(db: AsyncSession, project_id: uuid.UUID) -> Dict[str, Any]:
        stmt = text("""
            SELECT executive_summary, tender_type, complexity_level,
                   technical_requirements, commercial_requirements, legal_requirements,
                   required_documents, key_deadlines, risks
            FROM analysis_results
            WHERE project_id = :pid AND is_current = true AND status = 'completed'
        """)
        res = await db.execute(stmt, {"pid": project_id})
        row = res.mappings().first()
        return dict(row) if row else {}

    @staticmethod
    async def _load_chat_context(db: AsyncSession, project_id: uuid.UUID) -> Dict[str, Any]:
        stmt = text("""
            SELECT context FROM chat_sessions WHERE project_id = :pid
            ORDER BY updated_at DESC LIMIT 1
        """)
        res = await db.execute(stmt, {"pid": project_id})
        row = res.mappings().first()
        if not row or not row["context"]:
            return {}
        try:
            return row["context"] if isinstance(row["context"], dict) else json.loads(row["context"])
        except Exception:
            return {}

    @staticmethod
    def _load_company(company) -> Dict[str, Any]:
        return {
            "name": getattr(company, "name", ""),
            "bin_iin": getattr(company, "bin_iin", ""),
            "description": getattr(company, "description", ""),
            "specialization": getattr(company, "specialization", ""),
            "director_name": getattr(company, "director_name", ""),
            "director_title": getattr(company, "director_title", ""),
            "phone": getattr(company, "phone", ""),
            "email": getattr(company, "email", ""),
            "legal_address": getattr(company, "legal_address", ""),
        }

    @staticmethod
    def _compact_json(value) -> str:
        if value is None:
            return "—"
        if isinstance(value, str):
            try:
                parsed = json.loads(value)
                return json.dumps(parsed, ensure_ascii=False, indent=1)
            except Exception:
                return value
        return json.dumps(value, ensure_ascii=False, indent=1)

    @staticmethod
    async def generate(
        db: AsyncSession,
        project,
        company,
        req: GenerateRequest,
    ) -> GeneratedDocument:
        """Generate a tender document based on analysis + chat context + company profile."""
        analysis = await GenerationService._load_analysis_context(db, project.id)
        chat = await GenerationService._load_chat_context(db, project.id)
        company_data = GenerationService._load_company(company)

        # Idempotency: if the latest version of this doc type is already ready,
        # return it instead of spending another LLM call.
        latest = await GenerationService.get_latest(db, project.id, req.doc_type)
        if latest and latest.generation_status == "ready":
            logger.info("generation_reused_existing", project_id=str(project.id), doc_type=req.doc_type, version=latest.version)
            if project.status in ("generating", "clarifying", "analyzing", "draft"):
                project.status = "ready"
                await db.commit()
            return latest

        # Latest version
        version = (latest.version if latest else 0) + 1

        context_blocks = [
            "## Данные компании",
            json.dumps(company_data, ensure_ascii=False, indent=1),
            "## Анализ технического задания",
            json.dumps(
                {
                    k: GenerationService._compact_json(analysis.get(k))
                    for k in ["executive_summary", "tender_type", "complexity_level",
                              "technical_requirements", "commercial_requirements",
                              "legal_requirements", "required_documents", "key_deadlines", "risks"]
                },
                ensure_ascii=False,
                indent=1,
            ),
            "## Уточнённые данные из диалога (clarification context)",
            json.dumps(chat, ensure_ascii=False, indent=1),
        ]
        prompt = "\n\n".join(context_blocks) + f"\n\nЗадача:\n{DOC_TYPE_PROMPTS[req.doc_type]}"

        doc = GeneratedDocument(
            project_id=project.id,
            company_id=company.id,
            doc_type=req.doc_type,
            version=version,
            title=DOC_TYPE_LABELS.get(req.doc_type, req.doc_type),
            generation_status="generating",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        try:
            start = time.monotonic()
            parsed, model = await call_llm(
                prompt=prompt,
                system_prompt=GENERATION_SYSTEM_PROMPT,
                schema_class=GeneratedContent,
                estimated_tokens=2048,
            )
            elapsed_ms = int((time.monotonic() - start) * 1000)

            content_md = parsed.content_md or ""
            title = parsed.title or doc.title
            content_html = _full_html(title, _md_to_html(content_md))

            doc.title = title
            doc.content_md = content_md
            doc.content_html = content_html
            doc.generation_status = "ready"
            doc.llm_model = model
            doc.error_message = None
            if project.status in ("generating", "clarifying", "analyzing", "draft"):
                project.status = "ready"

            # Remove stale failed versions of the same doc_type to avoid clutter
            stale_stmt = (
                select(GeneratedDocument)
                .where(
                    GeneratedDocument.project_id == project.id,
                    GeneratedDocument.doc_type == req.doc_type,
                    GeneratedDocument.id != doc.id,
                    GeneratedDocument.generation_status == "failed",
                )
            )
            stale = (await db.execute(stale_stmt)).scalars().all()
            for s in stale:
                await db.delete(s)
        except Exception as e:
            logger.error("generation_failed", project_id=str(project.id), doc_type=req.doc_type, error=str(e))
            doc.generation_status = "failed"
            doc.error_message = str(e)[:1000]
            # Roll the project back so the user can retry generation (don't leave it stuck in "generating")
            if project.status == "generating":
                project.status = "clarifying"

        await db.commit()
        await db.refresh(doc)
        return doc

    @staticmethod
    async def list_generated(db: AsyncSession, project_id: uuid.UUID) -> List[GeneratedDocument]:
        stmt = select(GeneratedDocument).where(GeneratedDocument.project_id == project_id).order_by(
            GeneratedDocument.doc_type, GeneratedDocument.version
        )
        res = await db.execute(stmt)
        return list(res.scalars().all())

    @staticmethod
    async def get_latest(db: AsyncSession, project_id: uuid.UUID, doc_type: str) -> Optional[GeneratedDocument]:
        stmt = (
            select(GeneratedDocument)
            .where(GeneratedDocument.project_id == project_id, GeneratedDocument.doc_type == doc_type)
            .order_by(GeneratedDocument.version.desc())
            .limit(1)
        )
        res = await db.execute(stmt)
        return res.scalars().first()

    @staticmethod
    def _md_to_docx_bytes(title: str, md: str) -> bytes:
        """Render markdown (including tables) into a DOCX file using python-docx."""
        import io
        from docx import Document as DocxDocument
        from docx.shared import Pt

        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        doc.add_heading(title, level=0)

        def flush_list(list_buffer):
            for item in list_buffer:
                doc.add_paragraph(item, style="List Bullet")

        def add_table(rows):
            """rows: list of list[str] (first row = header)"""
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row):
                    table.cell(r_idx, c_idx).text = cell

        def add_blocks(text: str):
            list_buffer = []
            lines = text.splitlines()
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    if list_buffer:
                        flush_list(list_buffer)
                        list_buffer = []
                    continue

                # Markdown table: separator row like |---|---|
                if line.startswith("|") and i + 1 < len(lines):
                    sep = lines[i + 1].strip()
                    if re_fullmatch(r"\|?[\s:|-]+\|?[\s:|-]*", sep) and "|" in sep:
                        header = [c.strip() for c in line.strip("|").split("|")]
                        table_rows = [header]
                        i += 2
                        while i < len(lines) and lines[i].strip().startswith("|"):
                            row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                            table_rows.append(row)
                            i += 1
                        add_table(table_rows)
                        if list_buffer:
                            flush_list(list_buffer)
                            list_buffer = []
                        continue

                if line.startswith("### "):
                    flush_list(list_buffer); list_buffer = []
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    flush_list(list_buffer); list_buffer = []
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    flush_list(list_buffer); list_buffer = []
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    list_buffer.append(line[2:])
                else:
                    if list_buffer:
                        flush_list(list_buffer)
                        list_buffer = []
                    doc.add_paragraph(line)
                i += 1
            if list_buffer:
                flush_list(list_buffer)

        add_blocks(md)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    async def export(
        db: AsyncSession,
        project_id: uuid.UUID,
        doc_type: str,
        fmt: str,
    ) -> Tuple[bytes, str, str]:
        """Render the latest generated document as DOCX (python-docx) or PDF (Gotenberg LibreOffice)."""
        stmt = (
            select(GeneratedDocument)
            .where(GeneratedDocument.project_id == project_id, GeneratedDocument.doc_type == doc_type)
            .order_by(GeneratedDocument.version.desc())
            .limit(1)
        )
        res = await db.execute(stmt)
        doc = res.scalars().first()

        if not doc:
            raise HTTPException(status_code=404, detail="Сгенерированный документ не найден")
        if doc.generation_status != "ready" or not doc.content_md:
            raise HTTPException(status_code=409, detail="Документ ещё не готов к экспорту")

        if fmt == "docx":
            content = GenerationService._md_to_docx_bytes(doc.title, doc.content_md)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            url = settings.GOTENBERG_URL.rstrip("/") + "/forms/libreoffice/convert?format=pdf"
            files = {
                "files": (
                    f"{doc_type}.html",
                    doc.content_html.encode("utf-8"),
                    "text/html",
                )
            }
            async with httpx.AsyncClient(timeout=120) as client:
                resp = await client.post(url, files=files)
                if resp.status_code >= 400:
                    logger.error("gotenberg_export_failed", status=resp.status_code, body=resp.text[:500])
                    raise HTTPException(status_code=502, detail="Ошибка сервера экспорта")
            content = resp.content
            mime = "application/pdf"

        ext = fmt.lower()
        filename = f"{doc.title.replace(' ', '_')}_v{doc.version}.{ext}"

        exported = list(doc.exported_formats or [])
        if ext not in exported:
            exported.append(ext)
            doc.exported_formats = exported
            await db.commit()

        return content, filename, mime