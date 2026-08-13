from pypdf import PdfReader
from docx import Document
import re
from datetime import date
from typing import List, Dict, Any, Optional
import io

class DocumentParser:
    """
    Handles extraction, cleaning, and semantic chunking of text from PDF and DOCX files.
    """
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> str:
        """Extracts text from PDF bytes."""
        text_parts = []
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
            
        return "\n".join(text_parts)

    @staticmethod
    def get_page_count(file_bytes: bytes, mime_type: str) -> Optional[int]:
        """Best-effort page count for PDF files. Returns None for other types."""
        try:
            if mime_type == "application/pdf":
                reader = PdfReader(io.BytesIO(file_bytes))
                return len(reader.pages)
        except Exception:
            return None
        return None

    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> str:
        """Extracts text from DOCX bytes, including table cells."""
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

            # Tables often carry the actual tender requirements
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    text_parts.append("\n".join(rows))
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
            
        return "\n\n".join(text_parts)

    @staticmethod
    def clean_text(raw_text: str) -> str:
        """
        Cleans the extracted text:
        - Removes excessive newlines and spaces.
        - Joins artificially broken lines.
        - Removes simple headers/footers (basic regex).
        """
        if not raw_text:
            return ""
            
        # 1. Remove duplicate spaces
        text = re.sub(r'[ \t]+', ' ', raw_text)
        
        # 2. Normalize newlines (max 2 consecutive newlines)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 3. Join words broken by hyphens at the end of a line
        text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)
        
        # 4. Remove common standalone page numbers (e.g., lines with just a number)
        text = re.sub(r'(?m)^\s*\d+\s*$', '', text)
        
        return text.strip()

    @staticmethod
    def extract_metadata(cleaned_text: str) -> Dict[str, Any]:
        """
        Extracts document metadata (title, number, date) via regex heuristics.
        Returns a dict with optional keys: title, number, date.
        """
        meta: Dict[str, Any] = {}
        if not cleaned_text:
            return meta

        # 1. Document date: DD.MM.YYYY / DD-MM-YYYY / YYYY-MM-DD, optionally "от 12.03.2024"
        date_patterns = [
            re.compile(r'\b(\d{2})[.\-](\d{2})[.\-](\d{4})\b'),
            re.compile(r'\b(\d{4})[.\-](\d{2})[.\-](\d{2})\b'),
        ]
        for pattern in date_patterns:
            m = pattern.search(cleaned_text[:3000])
            if m:
                try:
                    if len(m.group(1)) == 4:
                        meta["date"] = date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
                    else:
                        meta["date"] = date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
                    break
                except ValueError:
                    continue

        # 2. Document number: "№ 123", "N 123", "номер 123", "№123-А", ISO-ish codes
        number_patterns = [
            re.compile(r'(?:№|N[оo]?\.?|номер)\s*([0-9]{1,6}(?:[-/][0-9A-Za-zА-Яа-яЁё]{1,10})*)', re.IGNORECASE),
            re.compile(r'\b(?:[A-Z]{2,5}-[0-9]{2,6}(?:[-/][0-9A-Za-zА-Яа-яЁё]{1,10})*)\b'),
        ]
        for pattern in number_patterns:
            m = pattern.search(cleaned_text[:3000])
            if m:
                meta["number"] = m.group(1) if len(m.groups()) else m.group(0)
                break

        # 3. Title: first non-empty line that is not a header/footer-ish fragment
        for line in cleaned_text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            # Skip lines that look like page numbers, dates, or short noise
            if len(stripped) < 4:
                continue
            if re.fullmatch(r'[\d\s./\-:]+', stripped):
                continue
            meta["title"] = stripped[:500]
            break

        return meta

    @staticmethod
    def detect_language(cleaned_text: str) -> Optional[str]:
        """Heuristic language detection: ru / kk / en based on character frequency."""
        if not cleaned_text:
            return None
        sample = cleaned_text[:5000]
        cyrillic = sum(1 for ch in sample if 'а' <= ch.lower() <= 'я')
        latin = sum(1 for ch in sample if 'a' <= ch.lower() <= 'z')
        if cyrillic == 0 and latin == 0:
            return None
        # Kazakh-specific letters distinguish kk from ru
        kz_letters = set("әіңғүұқөһӘІҢҒҮҰҚӨҺ")
        if cyrillic > latin and any(ch in kz_letters for ch in sample):
            return "kk"
        return "ru" if cyrillic > latin else "en"

    @staticmethod
    def chunk_text(cleaned_text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """
        Splits the text into semantic chunks using a sliding window.
        Returns a list of dictionaries with chunk metadata.
        """
        if not cleaned_text:
            return []
            
        chunks = []
        length = len(cleaned_text)
        start = 0
        chunk_index = 0
        
        while start < length:
            end = min(start + chunk_size, length)
            
            # Try to snap to the nearest newline or period if we are not at the end
            if end < length:
                # Find last newline within the chunk to avoid breaking paragraphs
                last_newline = cleaned_text.rfind('\n', start, end)
                if last_newline != -1 and (end - last_newline) < 300:
                    end = last_newline + 1
                else:
                    # Fallback to last period
                    last_period = cleaned_text.rfind('. ', start, end)
                    if last_period != -1 and (end - last_period) < 200:
                        end = last_period + 2
            
            chunk_text = cleaned_text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "index": chunk_index,
                    "text": chunk_text,
                    "char_length": len(chunk_text),
                    "start_idx": start,
                    "end_idx": end
                })
                chunk_index += 1
                
            start = end - overlap
            # Ensure we always move forward
            if start <= 0 or start == (end - overlap) and end == length:
                break
                
        return chunks
