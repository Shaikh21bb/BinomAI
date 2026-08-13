import pytest
import io

from app.services.generation_service import GenerationService, _md_to_html, _full_html

MD_WITH_TABLE = """## Спецификация элементов крыши

| № | Наименование | Характеристика | Ед. изм. | Кол-во |
|---|---|---|---|---|
| 1 | Металлочерепица | МП Супермонтеррей, 0.5 мм | м2 | 789,1 |
| 2 | Доска обрезная | 25х150 мм, сорт 1 | м3 | 8,95 |

- пункт списка один
- пункт списка два

Обычный абзац после таблицы.
"""


def test_md_to_html_renders_table():
    html = _md_to_html(MD_WITH_TABLE)
    assert "<table>" in html
    assert "Металлочерепица" in html
    assert "<th>" in html or "<td>" in html


def test_md_to_docx_renders_table():
    data = GenerationService._md_to_docx_bytes("Тест", MD_WITH_TABLE)

    from docx import Document as DocxDocument
    doc = DocxDocument(io.BytesIO(data))

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3  # header + 2 data rows
    assert len(table.columns) == 5
    assert table.cell(0, 1).text == "Наименование"
    assert table.cell(1, 1).text == "Металлочерепица"
    assert table.cell(2, 4).text == "8,95"

    # headings and paragraphs survive
    texts = [p.text for p in doc.paragraphs]
    assert any("Спецификация элементов крыши" in t for t in texts)
    assert any("Обычный абзац после таблицы." in t for t in texts)
